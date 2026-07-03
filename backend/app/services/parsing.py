"""Resume file parsing: PDF (PyMuPDF), DOCX (python-docx), plain text."""
import io

import fitz  # PyMuPDF
from docx import Document


def parse_resume(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _parse_pdf(content)
    elif name.endswith(".docx"):
        text = _parse_docx(content)
    elif name.endswith(".txt"):
        text = content.decode("utf-8", errors="replace")
    else:
        raise ValueError("Unsupported resume format — upload a .pdf, .docx, or .txt file.")
    text = text.strip()
    if len(text) < 50:
        raise ValueError("Could not extract text from the resume (is it a scanned image?).")
    return text


def _parse_pdf(content: bytes) -> str:
    with fitz.open(stream=content, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def _parse_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)
